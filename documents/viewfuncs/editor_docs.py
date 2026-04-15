# Editor Document Functions
# Contains create for Document model (for editor type)

import urllib.parse, requests, io, os, logging, platform, shutil, subprocess
from bs4 import BeautifulSoup
from ckeditor_uploader.views import upload as ckeditor_upload
from documents.forms import CreateDocumentForm
from documents.models import Folder, File, Document
from django.http import HttpResponseForbidden
from django.core.exceptions import ValidationError
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render, redirect
from django.utils.text import slugify
from docx import Document as DocxDocument
from docx.shared import Inches
from raadaa import settings
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)

# Formatting content in CKeditor
def add_formatted_content(doc, soup, word_dir):
    """Parse HTML and add formatted content and images to the .docx document."""
    current_paragraph = None

    def add_text_to_paragraph(text, bold=False, italic=False):
        """Helper to add text to the current paragraph with formatting."""
        nonlocal current_paragraph
        if not current_paragraph:
            current_paragraph = doc.add_paragraph()
        run = current_paragraph.add_run(text or '')
        run.bold = bold
        run.italic = italic

    for element in soup.find_all(recursive=False):  # Process top-level elements only
        if element.name == 'h1':
            current_paragraph = doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            current_paragraph = doc.add_heading(element.get_text(), level=2)
        elif element.name == 'p':
            current_paragraph = doc.add_paragraph()
            # Process nested elements within <p>
            for child in element.children:
                if child.name == 'strong':
                    add_text_to_paragraph(child.get_text(), bold=True)
                elif child.name == 'em':
                    add_text_to_paragraph(child.get_text(), italic=True)
                elif child.name == 'img':
                    img_src = child.get('src')
                    if img_src:
                        try:
                            parsed_src = urllib.parse.urlparse(img_src)
                            if parsed_src.scheme in ('http', 'https'):
                                print(f"Downloading remote image: {img_src}")
                                response = requests.get(img_src, timeout=5)
                                if response.status_code == 200:
                                    img_data = io.BytesIO(response.content)
                                    doc.add_picture(img_data, width=Inches(4.0))
                                else:
                                    print(f"Failed to download image: {img_src} (Status: {response.status_code})")
                                    current_paragraph = doc.add_paragraph(f"[Image failed to load: {img_src}]")
                            else:
                                # Handle images from CKEditor uploads
                                clean_src = img_src.replace(settings.MEDIA_URL, '').lstrip('/')

                                if settings.DEBUG:
                                    # Local dev: images exist on disk
                                    img_path = os.path.normpath(os.path.join(settings.MEDIA_ROOT, clean_src))
                                    print(f"Attempting to add local image: {img_path}")
                                    if os.path.exists(img_path):
                                        doc.add_picture(img_path, width=Inches(4.0))
                                    else:
                                        print(f"Local image not found: {img_path}")
                                        doc.add_paragraph(f"[Image not found: {img_path}]")
                                else:
                                    # Production: files are on S3 (remote)
                                    file_url = settings.MEDIA_URL + clean_src
                                    print(f"Attempting to fetch image from: {file_url}")
                                    try:
                                        response = requests.get(file_url)
                                        if response.status_code == 200:
                                            image_stream = io.BytesIO(response.content)
                                            doc.add_picture(image_stream, width=Inches(4.0))
                                        else:
                                            print(f"Image not accessible at: {file_url}")
                                            doc.add_paragraph(f"[Image not found: {file_url}]")
                                    except Exception as e:
                                        print(f"Error fetching image: {e}")
                                        doc.add_paragraph(f"[Error loading image: {file_url}]")
                        except Exception as e:
                            print(f"Error adding image {img_src}: {e}")
                            current_paragraph = doc.add_paragraph(f"[Error loading image: {img_src}]")
                else:
                    add_text_to_paragraph(child.get_text() if hasattr(child, 'get_text') else str(child))
        elif element.name == 'ul':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListBullet')
        elif element.name == 'ol':
            current_paragraph = None
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='ListNumber')
        else:
            current_paragraph = None
            add_text_to_paragraph(element.get_text() if hasattr(element, 'get_text') else str(element))

# Upload media in CKEditor helper function
@login_required
@csrf_exempt  # Required for CKEditor’s POST uploads
def custom_ckeditor_upload(request):
    if hasattr(request, 'tenant') and request.user.tenant != request.tenant:
        return HttpResponseForbidden("You are not authorized to perform actions for this company.")
    if not request.user.is_authenticated:
        return HttpResponseForbidden("You must be logged in to upload images.")
    logger.info(f"User {request.user.username} uploading image to CKEditor")
    print(f"User {request.user.username} uploading image to CKEditor")
    response = ckeditor_upload(request)
    if response.status_code == 200:
        logger.info(f"Image upload successful for user {request.user.username}")
        print(f"Image upload successful for user {request.user.username}")
    else:
        logger.error(f"Image upload failed for user {request.user.username}: {response.content}")
        print(f"Image upload failed for user {request.user.username}: {response.content}")
    return response

# Create editor document
@login_required
def create_from_editor(request):
    # Validate that the user belongs to the current tenant
    if request.user.tenant != request.tenant:
        return HttpResponseForbidden("Unauthorized: User does not belong to this company.")
    
    def documents_word_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'word')

    def documents_pdf_upload(request):
        tenant_name = request.tenant.name
        username = request.user.username if request.user else "anonymous"
        return os.path.join('documents', tenant_name, username, 'pdf')

    if request.method == "POST":
        form = CreateDocumentForm(request.POST)
        if form.is_valid():
            title = form.cleaned_data["title"]
            content = form.cleaned_data["content"]

            # Create a .docx file
            doc = DocxDocument()
            doc.add_heading(title, level=1)            # Parse HTML content using BeautifulSoup
            soup = BeautifulSoup(content, 'html.parser')

            # Add formatted content and images
            # Note: We pass None for word_dir because images will be handled via S3/URLs now
            add_formatted_content(doc, soup, None)

            # Get or create the "Template Document" folder
            user = request.user
            
            try:
                template_folder, created = Folder.objects.get_or_create(
                    tenant=request.tenant,
                    name="Template Document",
                    defaults={
                        'created_by': user,
                        'is_public': True,
                    }
                )
            except ValidationError as e:
                print(f"Folder creation error: {e}")
                messages.error(request, f"Error creating Template Document folder: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save the .docx file
            word_filename = f"{slugify(title)}_{request.user.id}_{template_folder.id}.docx"
            
            # Use io.BytesIO to save the document in memory first
            word_io = io.BytesIO()
            try:
                doc.save(word_io)
                word_io.seek(0)
                print(f"Generated .docx in memory: {word_filename}")
            except Exception as e:
                messages.error(request, f"Error creating .docx content: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Save to Document model first to use its file saving logic
            document = Document(
                document_type='Uploaded',
                document_source='editor',
                company_name='N/A',
                company_address='N/A',
                contact_person_name='N/A',
                contact_person_email='N/A',
                contact_person_designation='N/A',
                sales_rep='N/A',
                created_by=request.user,
                tenant=request.tenant,
            )
            
            try:
                # Save the word file using Django's FileField save method
                document.word_file.save(word_filename, ContentFile(word_io.read()), save=False)
                # Persist raw title and content for future editing
                document.editor_title = title
                document.editor_content = content
                document.save()
                print(f"Saved Document model with Word file: {document.word_file.url}")
            except Exception as e:
                messages.error(request, f"Error saving document record: {e}")
                return render(request, 'documents/create_from_editor.html', {'form': form})

            # Also save as a standalone File object in the "Template Document" folder
            try:
                # Reset word_io for a fresh read
                word_io.seek(0)
                public_word_file = File(
                    tenant=request.tenant,
                    original_name=word_filename,
                    folder=template_folder,
                    uploaded_by=user
                )
                public_word_file.file.save(word_filename, ContentFile(word_io.read()), save=True)
                print(f"Saved standalone File: {public_word_file.file.url}")
            except Exception as e:
                print(f"Warning: Error saving to File model: {e}")
                # We don't return here because the main Document was saved

            messages.success(request, 'Document created successfully!')
            return redirect("document_list")
        else:
            print("Form errors:", form.errors)
            messages.error(request, 'Please correct the errors below.')
    else:
        form = CreateDocumentForm()
    return render(request, 'documents/create_from_editor.html', {'form': form})

@login_required
def generate_pdf_view(request, document_id):
    """Manually trigger PDF generation for a document."""
    from django.shortcuts import get_object_or_404
    document = get_object_or_404(Document, id=document_id, tenant=request.tenant)
    
    if not document.word_file:
        messages.error(request, "Word file not found for this document.")
        return redirect("document_list")

    # If PDF already exists, maybe we shouldn't overwrite? Or we do.
    # For now, let's allow regenerating.

    # 1. Download/Get the word file content.
    # We explicitly open() before read() because django-storages (S3)
    # returns a lazy file descriptor that must be opened first.
    # For local storage this is a no-op.
    try:
        document.word_file.open('rb')
        word_content = document.word_file.read()
        document.word_file.close()
    except Exception as e:
        messages.error(request, f"Could not read Word file: {e}")
        return redirect("document_list")

    # 2. Temporary storage for conversion
    import tempfile
    with tempfile.TemporaryDirectory() as temp_dir:
        # Use os.path.basename to strip S3 key prefixes from the stored path
        # e.g. "documents/TenantName/user/word/file.docx" -> "file.docx"
        word_basename = os.path.basename(document.word_file.name)
        word_temp_path = os.path.join(temp_dir, word_basename)
        with open(word_temp_path, 'wb') as f:
            f.write(word_content)
        logger.info(f"Word file written to temp path: {word_temp_path}")

        # 3. Determine LibreOffice path
        if platform.system() == "Windows":
            paths = [
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
                r"C:\Program Files\LibreOffice\program\soffice.exe"
            ]
            libreoffice_path = next((p for p in paths if os.path.exists(p)), None)
        else:
            libreoffice_path = shutil.which("libreoffice") or shutil.which("soffice")

        if not libreoffice_path:
            messages.error(request, "LibreOffice not found on the server. Please contact admin.")
            return redirect("document_list")

        # 4. Run conversion
        try:
            print(f"Starting PDF conversion for document {document_id}")
            subprocess.run(
                [libreoffice_path, "--headless", "--convert-to", "pdf", "--outdir", temp_dir, word_temp_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                timeout=30
            )
            
            # The output file will have the same basename but .pdf extension
            pdf_filename = os.path.splitext(word_basename)[0] + ".pdf"
            pdf_temp_path = os.path.join(temp_dir, pdf_filename)

            if os.path.exists(pdf_temp_path):
                with open(pdf_temp_path, 'rb') as f:
                    pdf_bytes = f.read()

                # Save PDF — django-storages automatically uploads to S3 in
                # production or writes to MEDIA_ROOT locally, depending on
                # DEFAULT_FILE_STORAGE / STORAGES settings.
                document.pdf_file.save(pdf_filename, ContentFile(pdf_bytes), save=True)
                logger.info(f"PDF saved to storage: {document.pdf_file.name}")
                messages.success(request, f"PDF generated successfully for {word_basename}")
            else:
                messages.error(request, "LibreOffice failed to generate PDF.")
        except Exception as e:
            print(f"PDF conversion error: {e}")
            logger.error(f"PDF conversion error for document {document_id}: {e}")
            messages.error(request, f"Error during PDF conversion: {e}")

    return redirect("document_list")


@login_required
def edit_editor_document(request, document_id):
    """Edit the title and content of an editor-created document.
    
    Only the document owner (created_by) or users with the Admin role can edit.
    On save the .docx is regenerated and the existing PDF is cleared so it
    can be regenerated on demand.
    """
    from django.shortcuts import get_object_or_404

    document = get_object_or_404(
        Document,
        id=document_id,
        tenant=request.tenant,
        document_source='editor',
    )

    # --- Permission check ---
    is_owner = document.created_by == request.user
    is_admin = request.user.roles.filter(name='Admin').exists()
    if not (is_owner or is_admin):
        return HttpResponseForbidden("You do not have permission to edit this document.")

    if request.method == "POST":
        form = CreateDocumentForm(request.POST)
        if form.is_valid():
            title = form.cleaned_data["title"]
            content = form.cleaned_data["content"]

            # Re-build the .docx from the updated content
            doc = DocxDocument()
            doc.add_heading(title, level=1)
            soup = BeautifulSoup(content, 'html.parser')
            add_formatted_content(doc, soup, None)

            word_filename = f"{slugify(title)}_{request.user.id}_{document.id}.docx"
            word_io = io.BytesIO()
            try:
                doc.save(word_io)
                word_io.seek(0)
            except Exception as e:
                messages.error(request, f"Error rebuilding .docx: {e}")
                return render(request, 'documents/edit_editor_document.html', {'form': form, 'document': document})

            # Delete old word file from storage before saving the new one
            if document.word_file:
                try:
                    document.word_file.delete(save=False)
                except Exception:
                    pass  # File may not exist on storage yet

            # Delete old PDF so user is prompted to regenerate
            if document.pdf_file:
                try:
                    document.pdf_file.delete(save=False)
                    document.pdf_file = None
                except Exception:
                    document.pdf_file = None

            # Save the new .docx and persist updated title/content
            try:
                document.word_file.save(word_filename, ContentFile(word_io.read()), save=False)
                document.editor_title = title
                document.editor_content = content
                document.save()
                logger.info(f"Document {document_id} updated by {request.user.username}")
            except Exception as e:
                messages.error(request, f"Error saving updated document: {e}")
                return render(request, 'documents/edit_editor_document.html', {'form': form, 'document': document})

            messages.success(request, "Document updated successfully! PDF has been cleared — click 'Generate PDF' to create a fresh one.")
            return redirect("document_list")
        else:
            messages.error(request, "Please correct the errors below.")
    else:
        # Pre-fill with stored title and content
        initial = {
            'title': document.editor_title or '',
            'content': document.editor_content or '',
        }
        form = CreateDocumentForm(initial=initial)

    return render(request, 'documents/edit_editor_document.html', {'form': form, 'document': document})